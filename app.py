import os
import io
import re
import json
from datetime import datetime
from typing import Dict, List, Tuple

import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
from openai import OpenAI

from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader


# =========================
# Config
# =========================
TEMPLATE_XLSX = "marronisme_onderzoek_template.xlsx"

DATA_DIR = "data"
os.makedirs(DATA_DIR, exist_ok=True)
OUT_CSV = os.path.join(DATA_DIR, "responses_formulier.csv")

# fallback; we override with dimensions sheet if present
DIM_ORDER_DEFAULT = ["historisch", "sociaal", "moreel", "spiritueel", "praktisch"]

# OpenAI client (leest OPENAI_API_KEY automatisch uit env / Streamlit Secrets)
client = OpenAI()


# =========================
# Utilities
# =========================
def safe_lower(x: str) -> str:
    return str(x or "").strip().lower()


def word_count(s: str) -> int:
    s = str(s or "")
    # eenvoudige woordtelling (spaties + interpunctie)
    words = re.findall(r"\b\w+\b", s, flags=re.UNICODE)
    return len(words)


def char_count(s: str) -> int:
    return len(str(s or ""))


def get_query_param(name: str) -> str:
    """
    Compatibel met meerdere Streamlit versies.
    """
    # Nieuwere Streamlit heeft st.query_params
    try:
        qp = st.query_params
        if name in qp:
            v = qp.get(name)
            if isinstance(v, list):
                return v[0] if v else ""
            return str(v)
        return ""
    except Exception:
        pass

    # Oudere Streamlit
    try:
        qp = st.experimental_get_query_params()
        v = qp.get(name, [""])
        return v[0] if v else ""
    except Exception:
        return ""


# =========================
# Template + Opslag
# =========================
def load_dimensions() -> Tuple[List[str], Dict[str, str]]:
    """
    Returns:
      dim_order: list of dimensies in officiële volgorde
      dim_desc: dict dimensie -> beschrijving
    """
    try:
        dim_df = pd.read_excel(TEMPLATE_XLSX, sheet_name="dimensions", dtype=str).fillna("")
        dim_df["dimensie"] = dim_df["dimensie"].apply(safe_lower)
        dim_df["beschrijving"] = dim_df["beschrijving"].astype(str).fillna("").str.strip()
        dim_order = [d for d in dim_df["dimensie"].tolist() if d]
        dim_desc = {r["dimensie"]: r["beschrijving"] for _, r in dim_df.iterrows()}
        if dim_order:
            return dim_order, dim_desc
    except Exception:
        pass

    return DIM_ORDER_DEFAULT, {d: "" for d in DIM_ORDER_DEFAULT}


def load_template():
    """
    Expected columns in sheet 'questions':
      - vraag_id
      - dimensie
      - type (mc/open)
      - vraagtekst
      - theoretisch_kader
    Sheet 'mc_options':
      - vraag_id, optie, label
    """
    q = pd.read_excel(TEMPLATE_XLSX, sheet_name="questions", dtype=str).fillna("")
    m = pd.read_excel(TEMPLATE_XLSX, sheet_name="mc_options", dtype=str).fillna("")

    q["vraag_id"] = q["vraag_id"].astype(str).str.strip()
    q["dimensie"] = q["dimensie"].apply(safe_lower)
    q["type"] = q["type"].apply(safe_lower)
    q["vraagtekst"] = q.get("vraagtekst", "").astype(str).fillna("").str.strip()
    q["theoretisch_kader"] = q.get("theoretisch_kader", "").astype(str).fillna("").str.strip()

    # MC opties
    m["vraag_id"] = m["vraag_id"].astype(str).str.strip()
    m["optie"] = m["optie"].astype(str).str.strip().str.upper()
    m["label"] = m["label"].astype(str).str.strip()

    mc_map = {}
    for vid, sub in m.groupby("vraag_id"):
        order = {"A": 1, "B": 2, "C": 3, "D": 4}
        sub = sub.copy()
        sub["rank"] = sub["optie"].map(order).fillna(99)
        sub = sub.sort_values("rank")
        mc_map[vid] = [f"{row.optie} — {row.label}" for row in sub.itertuples(index=False)]

    # sorteer vragen binnen dimensies
    def qnum(v):
        try:
            return int(str(v).replace("Q", "").strip())
        except Exception:
            return 9999

    dim_order, _ = load_dimensions()
    dim_rank = {d: i for i, d in enumerate(dim_order)}

    q["qnum"] = q["vraag_id"].apply(qnum)
    q["dim_rank"] = q["dimensie"].map(dim_rank).fillna(999)
    q = q.sort_values(["dim_rank", "qnum"]).drop(columns=["qnum", "dim_rank"])

    # handige mapping voor later (vraagtekst + kader)
    q_map = {}
    for r in q.itertuples(index=False):
        q_map[str(r.vraag_id)] = {
            "dimensie": safe_lower(getattr(r, "dimensie", "")),
            "type": safe_lower(getattr(r, "type", "")),
            "vraagtekst": str(getattr(r, "vraagtekst", "")).strip(),
            "theoretisch_kader": str(getattr(r, "theoretisch_kader", "")).strip(),
        }

    return q, mc_map, q_map


def append_rows(rows):
    df = pd.DataFrame(rows)[["respondent_id", "vraag_id", "dimensie", "type", "antwoord", "timestamp"]]
    try:
        if os.path.exists(OUT_CSV):
            df.to_csv(OUT_CSV, mode="a", header=False, index=False, encoding="utf-8-sig")
        else:
            df.to_csv(OUT_CSV, index=False, encoding="utf-8-sig")
    except PermissionError:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        alt = os.path.join(DATA_DIR, f"responses_formulier_{ts}.csv")
        df.to_csv(alt, index=False, encoding="utf-8-sig")


def load_responses():
    if not os.path.exists(OUT_CSV):
        return pd.DataFrame(columns=["respondent_id", "vraag_id", "dimensie", "type", "antwoord", "timestamp"])
    return pd.read_csv(OUT_CSV, dtype=str).fillna("")


# =========================
# Privacy: namen eruit, locaties blijven
# =========================
def redact_names_keep_locations(text: str, names_to_redact: List[str]) -> str:
    """
    Verwijder opgegeven namen + basis-PII (email/telefoon). Locaties blijven staan.
    """
    if not text:
        return ""
    t = str(text)

    # Emails
    t = re.sub(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", "[EMAIL]", t)

    # Telefoonnummers
    t = re.sub(r"\b(\+?\d[\d\s().-]{7,}\d)\b", "[TEL]", t)

    # Exact opgegeven namen (case-insensitive)
    for name in names_to_redact:
        n = name.strip()
        if not n:
            continue
        t = re.sub(rf"\b{re.escape(n)}\b", "[NAAM]", t, flags=re.IGNORECASE)

    return t


# =========================
# Analyse: intensiteit + grafiek
# =========================
def intensity_by_dimension(df: pd.DataFrame, dim_order: List[str]) -> pd.DataFrame:
    rows = []
    for d in dim_order:
        sub = df[df["dimensie"] == d]
        text_all = "\n".join(sub["antwoord"].astype(str).tolist())
        rows.append(
            {
                "dimensie": d,
                "aantal_antwoorden": int(len(sub)),
                "woorden": word_count(text_all),
                "tekens": char_count(text_all),
            }
        )
    out = pd.DataFrame(rows)
    out["woorden_pct"] = (out["woorden"] / max(out["woorden"].sum(), 1) * 100).round(1)
    out["tekens_pct"] = (out["tekens"] / max(out["tekens"].sum(), 1) * 100).round(1)
    return out


def make_dimension_bar_chart_png(intensity_df: pd.DataFrame, metric: str = "woorden") -> bytes:
    """
    metric: 'woorden' or 'tekens'
    Returns PNG bytes.
    """
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.bar(intensity_df["dimensie"], intensity_df[metric])
    ax.set_title(f"Intensiteit per dimensie ({metric})")
    ax.set_xlabel("Dimensie")
    ax.set_ylabel(metric.capitalize())
    ax.tick_params(axis="x", rotation=25)

    bio = io.BytesIO()
    fig.tight_layout()
    fig.savefig(bio, format="png", dpi=180)
    plt.close(fig)
    return bio.getvalue()


# =========================
# AI Context: antwoorden + theoretisch kader
# =========================
def build_theoretical_framework(q_map: Dict[str, Dict[str, str]], dim_order: List[str]) -> Dict[str, List[str]]:
    """
    returns: dict dimensie -> list of unique theoretisch_kader items
    """
    by_dim = {d: [] for d in dim_order}
    for vid, meta in q_map.items():
        d = safe_lower(meta.get("dimensie", ""))
        tk = str(meta.get("theoretisch_kader", "")).strip()
        if d in by_dim and tk:
            by_dim[d].append(tk)

    # unique preserve order
    out = {}
    for d in dim_order:
        seen = set()
        uniq = []
        for x in by_dim.get(d, []):
            key = x.strip().lower()
            if key and key not in seen:
                seen.add(key)
                uniq.append(x.strip())
        out[d] = uniq
    return out


def df_to_ai_context(df: pd.DataFrame, q_map: Dict[str, Dict[str, str]], dim_order: List[str]) -> str:
    """
    Compacte context: per antwoord voeg vraagtekst + theoretisch kader toe.
    """
    lines = []
    # sorteer op dim_order, daarna vraag_id
    dim_rank = {d: i for i, d in enumerate(dim_order)}
    df2 = df.copy()
    df2["dim_rank"] = df2["dimensie"].map(dim_rank).fillna(999)
    df2 = df2.sort_values(["dim_rank", "vraag_id"]).drop(columns=["dim_rank"])

    for r in df2.itertuples(index=False):
        vid = str(r.vraag_id)
        meta = q_map.get(vid, {})
        vtxt = str(meta.get("vraagtekst", "")).strip()
        tk = str(meta.get("theoretisch_kader", "")).strip()
        ans = str(r.antwoord).strip()
        dim = safe_lower(r.dimensie)
        chunk = f"[{dim} | {vid}]"
        if vtxt:
            chunk += f" VRAAG: {vtxt}"
        if tk:
            chunk += f" | KADER: {tk}"
        chunk += f"\nANTWOORD: {ans}"
        lines.append(chunk)
    return "\n\n".join(lines)


# =========================
# AI: dimensie-tabel (JSON)
# =========================
def ai_dimensie_tabel(df_ai: pd.DataFrame, dim_order: List[str], dim_desc: Dict[str, str]) -> pd.DataFrame:
    context = df_to_ai_context(df_ai, st.session_state["q_map"], dim_order)

    prompt = f"""
Je bent een onderzoek-assistent.

Maak een samenvatting per dimensie op basis van onderzoeksantwoorden over Marronisme.

DIMENSIES (gebruik EXACT deze labels en deze volgorde):
{dim_order}

DIMENSIE-BESCHRIJVINGEN:
{json.dumps(dim_desc, ensure_ascii=False)}

EISEN:
- Output moet STRICT JSON zijn (geen markdown).
- Schema:
{{
  "rows": [
    {{"dimensie":"historisch","kerninzichten":["...","...","..."]}},
    ...
  ]
}}
- 3 t/m 6 kerninzichten per dimensie.
- Formuleer analytisch, academisch, en compact.
- Als er weinig data is: benoem dat, maar geef toch interpretatieve observaties op basis van context.

DATA:
{context}
""".strip()

    resp = client.responses.create(model=st.session_state["model_name"], input=prompt)
    raw = resp.output_text.strip()

    data = json.loads(raw)
    rows = data.get("rows", [])

    ai_map = {}
    for r in rows:
        d = safe_lower(r.get("dimensie", ""))
        bullets = r.get("kerninzichten", [])
        if isinstance(bullets, list):
            bullets = [str(b).strip() for b in bullets if str(b).strip()]
        else:
            bullets = []
        ai_map[d] = bullets

    final = []
    for d in dim_order:
        bullets = ai_map.get(d, [])
        final.append(
            {
                "dimensie": d,
                "kerninzichten": " • " + " • ".join(bullets) if bullets else "Geen expliciete data — interpretatieve observatie op basis van context.",
            }
        )
    return pd.DataFrame(final)


# =========================
# AI: APA-rapport
# =========================
def ai_apa_report(
    df_ai: pd.DataFrame,
    dim_table: pd.DataFrame,
    intensity_df: pd.DataFrame,
    dim_order: List[str],
    dim_desc: Dict[str, str],
    theoretical_framework_by_dim: Dict[str, List[str]],
) -> str:
    context = df_to_ai_context(df_ai, st.session_state["q_map"], dim_order)

    # compact helpers
    tf_compact = {d: theoretical_framework_by_dim.get(d, []) for d in dim_order}
    intensity_compact = intensity_df.to_dict(orient="records")
    dim_table_compact = dim_table.to_dict(orient="records")

    prompt = f"""
Schrijf een subsidie-proof academisch rapport in APA-stijl over de onderzoeksantwoorden m.b.t. Marronisme.

STRUCTUUR (APA, met kopjes):
1. Titel
2. Abstract (max 200 woorden)
3. Inleiding
4. Theoretisch Kader (gebruik de opgegeven "theoretisch_kader" thema’s; koppel ze aan de dimensies)
5. Methode (beschrijf dat het een vragenlijst is met dimensies; geen verzonnen steekproefgroottes; gebruik alleen data die gegeven is)
6. Resultaten
   - Verwerk de DIMENSIE-SAMENVATTING en de INTENSITEITSANALYSE in tekst
7. Discussie
8. Conclusie
9. Implicaties voor beleid/programma’s (subsidiegericht)
10. Beperkingen & Aanbevelingen voor vervolgonderzoek
11. Werkdefinitie van Marronisme (1 alinea + 3 bullets met kerncriteria)

REGELS:
- Wees academisch, helder, en niet “te literair”.
- Verzín geen feiten (geen aantallen respondenten behalve wat zichtbaar is).
- Verwijs expliciet naar de 5 dimensies en hun betekenissen.
- Gebruik neutrale toon.
- Houd het rapport bruikbaar voor subsidieaanvragen.

INPUTS:
DIMENSIES (volgorde):
{dim_order}

DIMENSIE-BESCHRIJVINGEN:
{json.dumps(dim_desc, ensure_ascii=False)}

THEORETISCH KADER PER DIMENSIE (thema’s uit template):
{json.dumps(tf_compact, ensure_ascii=False)}

DIMENSIE-SAMENVATTING (tabel):
{json.dumps(dim_table_compact, ensure_ascii=False)}

INTENSITEITSANALYSE (woorden/tekens):
{json.dumps(intensity_compact, ensure_ascii=False)}

DATA (vraag+antwoord+theoretisch kader):
{context}
""".strip()

    resp = client.responses.create(model=st.session_state["model_name"], input=prompt)
    return resp.output_text


# =========================
# Export: Word + PDF (met grafiek)
# =========================
def build_docx_bytes(
    title: str,
    respondent_id: str,
    dim_table: pd.DataFrame,
    intensity_df: pd.DataFrame,
    report_text: str,
    chart_png_bytes: bytes,
) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    meta = doc.add_paragraph(
        f"Respondent ID: {respondent_id}\nDatum/tijd: {datetime.now().isoformat(timespec='seconds')}"
    )
    for run in meta.runs:
        run.font.size = Pt(10)

    doc.add_heading("Samenvatting per dimensie", level=2)
    t1 = doc.add_table(rows=1, cols=2)
    hdr = t1.rows[0].cells
    hdr[0].text = "Dimensie"
    hdr[1].text = "Kerninzichten"
    for _, r in dim_table.iterrows():
        row = t1.add_row().cells
        row[0].text = str(r.get("dimensie", ""))
        row[1].text = str(r.get("kerninzichten", ""))

    doc.add_paragraph("")

    doc.add_heading("Intensiteitsanalyse", level=2)
    t2 = doc.add_table(rows=1, cols=5)
    hdr2 = t2.rows[0].cells
    hdr2[0].text = "Dimensie"
    hdr2[1].text = "# Antwoorden"
    hdr2[2].text = "Woorden"
    hdr2[3].text = "Tekens"
    hdr2[4].text = "Woorden (%)"
    for _, r in intensity_df.iterrows():
        row = t2.add_row().cells
        row[0].text = str(r.get("dimensie", ""))
        row[1].text = str(r.get("aantal_antwoorden", ""))
        row[2].text = str(r.get("woorden", ""))
        row[3].text = str(r.get("tekens", ""))
        row[4].text = str(r.get("woorden_pct", ""))

    doc.add_paragraph("")

    doc.add_heading("Grafiek", level=2)
    img_stream = io.BytesIO(chart_png_bytes)
    doc.add_picture(img_stream, width=Inches(6.5))

    doc.add_paragraph("")
    doc.add_heading("APA-rapport", level=2)
    # tekst in paragrafen splitsen
    for para in report_text.split("\n"):
        p = doc.add_paragraph(para)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_pdf_bytes(
    title: str,
    respondent_id: str,
    dim_table: pd.DataFrame,
    intensity_df: pd.DataFrame,
    report_text: str,
    chart_png_bytes: bytes,
) -> bytes:
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    w, h = A4
    x = 2 * cm
    y = h - 2 * cm

    def line(txt, dy=14, font="Helvetica", size=10):
        nonlocal y
        if y < 2 * cm:
            c.showPage()
            y = h - 2 * cm
        c.setFont(font, size)
        c.drawString(x, y, txt[:115])
        y -= dy

    line(title, dy=18, font="Helvetica-Bold", size=14)
    line(f"Respondent ID: {respondent_id}")
    line(f"Datum/tijd: {datetime.now().isoformat(timespec='seconds')}")
    line("")

    line("Samenvatting per dimensie", dy=16, font="Helvetica-Bold", size=12)
    for _, r in dim_table.iterrows():
        line(f"- {r.get('dimensie','')}:")
        txt = str(r.get("kerninzichten", "")).replace("•", "\n•")
        for b in txt.splitlines():
            b = b.strip()
            if b:
                line(f"  {b}")
        line("")

    line("Intensiteitsanalyse", dy=16, font="Helvetica-Bold", size=12)
    for _, r in intensity_df.iterrows():
        line(f"- {r.get('dimensie','')}: woorden={r.get('woorden','')}, tekens={r.get('tekens','')}, antwoorden={r.get('aantal_antwoorden','')}")
    line("")

    line("Grafiek", dy=16, font="Helvetica-Bold", size=12)
    # plaats afbeelding (als er ruimte is, anders nieuwe pagina)
    if y < 10 * cm:
        c.showPage()
        y = h - 2 * cm
    img = ImageReader(io.BytesIO(chart_png_bytes))
    c.drawImage(img, x, y - 8 * cm, width=16 * cm, height=8 * cm, preserveAspectRatio=True, mask='auto')
    y -= 9 * cm

    line("APA-rapport", dy=16, font="Helvetica-Bold", size=12)
    for para in report_text.splitlines():
        t = para.strip()
        if not t:
            line("")
            continue
        while len(t) > 110:
            line(t[:110])
            t = t[110:]
        line(t)

    c.save()
    return bio.getvalue()


# =========================
# App
# =========================
def main():
    st.set_page_config(page_title="Marronisme Onderzoek", layout="wide")
    st.title("Marronisme Onderzoek – Enquête + AI + Export")
    st.caption("Vragen uit Excel-template. Antwoorden naar CSV. AI-analyse met APA-rapport + grafiek + Word/PDF export.")

    if not os.path.exists(TEMPLATE_XLSX):
        st.error(f"Template niet gevonden: {TEMPLATE_XLSX}")
        st.stop()

    # model (aanpasbaar)
    with st.sidebar:
        st.subheader("Instellingen")
        st.session_state["model_name"] = st.text_input("OpenAI model", value=st.session_state.get("model_name", "gpt-5-mini"))
        st.markdown("---")
        st.subheader("Respondentenlink")
        app_url = st.text_input(
            "Plak hier jouw online Streamlit app-URL",
            value=st.session_state.get("app_url", ""),
            placeholder="https://jouwapp.streamlit.app",
        )
        st.session_state["app_url"] = app_url

    dim_order, dim_desc = load_dimensions()
    questions, mc_map, q_map = load_template()
    st.session_state["q_map"] = q_map

    tab1, tab2 = st.tabs(["Enquête invullen (respondent)", "AI analyse & export (onderzoeker)"])

    # -------------------------
    # TAB 1: Respondent invullen
    # -------------------------
    with tab1:
        st.subheader("Enquête invullen")
        st.write("Totaal vragen geladen:", len(questions))

        # Prefill rid via query param: ?rid=001
        rid_from_url = get_query_param("rid")
        respondent_id = st.text_input("Respondent ID (bijv. 001)", value=rid_from_url or "")

        # Toon respondent-link generator (handig voor jou als beheerder)
        if st.session_state.get("app_url"):
            st.info("Link om naar respondenten te sturen (vult Respondent ID automatisch):")
            example_rid = respondent_id.strip() if respondent_id.strip() else "001"
            st.code(f"{st.session_state['app_url'].rstrip('/')}/?rid={example_rid}")

        with st.form("form_all"):
            answers = []

            for dim in dim_order:
                sub = questions[questions["dimensie"] == dim]
                with st.expander(f"{dim.capitalize()} ({len(sub)} vragen)", expanded=(dim == dim_order[0])):
                    for _, q in sub.iterrows():
                        vid = q["vraag_id"]
                        qtype = q["type"]
                        vtext = str(q.get("vraagtekst", "")).strip()

                        st.markdown(f"### {vid}")
                        if vtext:
                            st.write(vtext)

                        if qtype == "mc":
                            options = mc_map.get(vid, ["A", "B", "C", "D"])
                            choice = st.radio("Kies één:", options, key=f"ans_{vid}")
                            answer = str(choice).split("—")[0].strip()
                        else:
                            answer = st.text_area("Antwoord:", key=f"ans_{vid}", height=120)

                        answers.append((vid, dim, qtype, answer))
                        st.divider()

            submitted = st.form_submit_button("Alles opslaan", type="primary")

            if submitted:
                if not respondent_id.strip():
                    st.error("Respondent ID is verplicht.")
                    st.stop()

                ts = datetime.now().isoformat(timespec="seconds")
                rows = []
                for vid, dim, t, ans in answers:
                    rows.append(
                        {
                            "respondent_id": respondent_id.strip(),
                            "vraag_id": vid,
                            "dimensie": dim,
                            "type": t,
                            "antwoord": (ans or "").strip(),
                            "timestamp": ts,
                        }
                    )

                append_rows(rows)
                st.success(f"Opgeslagen! → {OUT_CSV}")

    # -------------------------
    # TAB 2: Onderzoeker analyse & export
    # -------------------------
    with tab2:
        st.subheader("AI analyse & export")
        df_all = load_responses()

        if df_all.empty:
            st.info("Nog geen opgeslagen antwoorden gevonden. Laat eerst respondenten invullen.")
            return

        respondent_ids = sorted([x for x in df_all["respondent_id"].unique() if str(x).strip()])
        rid = st.selectbox("Kies respondent_id", respondent_ids)

        df_r = df_all[df_all["respondent_id"] == rid].copy()
        # verzeker dimensie labels lower
        df_r["dimensie"] = df_r["dimensie"].apply(safe_lower)

        st.markdown("### Privacy: namen eruit (locaties blijven)")
        names_raw = st.text_area(
            "Vul namen in die je wilt verwijderen (1 per regel). Locaties blijven staan.",
            height=120,
            placeholder="Bijv:\nKofi\nAma\nJan de Vries",
        )
        names_to_redact = [n.strip() for n in names_raw.splitlines() if n.strip()]

        # Redactie vóór AI
        df_ai = df_r.copy()
        df_ai["antwoord"] = df_ai["antwoord"].apply(lambda t: redact_names_keep_locations(t, names_to_redact))

        with st.expander("Bekijk geredigeerde antwoorden (dit gaat naar AI)", expanded=False):
            st.dataframe(df_ai[["vraag_id", "dimensie", "antwoord"]], use_container_width=True)

        # Intensiteitsanalyse + grafiek (altijd beschikbaar)
        intensity_df = intensity_by_dimension(df_ai, dim_order)
        st.markdown("### Intensiteitsanalyse (woorden/tekens per dimensie)")
        st.dataframe(intensity_df, use_container_width=True)

        chart_metric = st.selectbox("Grafiek-metriek", ["woorden", "tekens"], index=0)
        chart_png = make_dimension_bar_chart_png(intensity_df, metric=chart_metric)
        st.image(chart_png, caption=f"Intensiteit per dimensie ({chart_metric})", use_column_width=True)

        # AI knoppen
        c1, c2 = st.columns(2)

        with c1:
            if st.button("Maak AI dimensie-tabel", type="primary"):
                with st.spinner("AI maakt dimensie-samenvatting..."):
                    try:
                        dim_table = ai_dimensie_tabel(df_ai, dim_order, dim_desc)
                        st.session_state["dim_table"] = dim_table
                        st.dataframe(dim_table, use_container_width=True)
                        st.success("Dimensie-tabel gereed.")
                    except Exception as e:
                        st.error(f"AI fout (dimensie-tabel): {e}")

        with c2:
            if st.button("Maak AI APA-rapport"):
                if "dim_table" not in st.session_state:
                    st.warning("Maak eerst de AI dimensie-tabel (knop links).")
                else:
                    with st.spinner("AI schrijft APA-rapport..."):
                        try:
                            tf_by_dim = build_theoretical_framework(q_map, dim_order)
                            report = ai_apa_report(
                                df_ai=df_ai,
                                dim_table=st.session_state["dim_table"],
                                intensity_df=intensity_df,
                                dim_order=dim_order,
                                dim_desc=dim_desc,
                                theoretical_framework_by_dim=tf_by_dim,
                            )
                            st.session_state["apa_report"] = report
                            st.markdown(report)
                            st.success("APA-rapport gereed.")
                        except Exception as e:
                            st.error(f"AI fout (APA-rapport): {e}")

        st.markdown("---")
        st.markdown("## Export (Word / PDF)")

        dim_table = st.session_state.get("dim_table")
        apa_report = st.session_state.get("apa_report")

        if dim_table is None or apa_report is None:
            st.info("Maak eerst de AI dimensie-tabel én het AI APA-rapport om te exporteren.")
        else:
            try:
                docx_bytes = build_docx_bytes(
                    title="Marronisme Onderzoek – AI Rapport (APA)",
                    respondent_id=rid,
                    dim_table=dim_table,
                    intensity_df=intensity_df,
                    report_text=apa_report,
                    chart_png_bytes=chart_png,
                )
                pdf_bytes = build_pdf_bytes(
                    title="Marronisme Onderzoek – AI Rapport (APA)",
                    respondent_id=rid,
                    dim_table=dim_table,
                    intensity_df=intensity_df,
                    report_text=apa_report,
                    chart_png_bytes=chart_png,
                )

                d1, d2 = st.columns(2)
                with d1:
                    st.download_button(
                        "Download Word (.docx)",
                        data=docx_bytes,
                        file_name=f"marronisme_ai_apa_rapport_{rid}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                with d2:
                    st.download_button(
                        "Download PDF (.pdf)",
                        data=pdf_bytes,
                        file_name=f"marronisme_ai_apa_rapport_{rid}.pdf",
                        mime="application/pdf",
                    )
            except Exception as e:
                st.error(f"Export fout: {e}")


if __name__ == "__main__":
    main()